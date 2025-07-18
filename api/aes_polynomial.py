# api/aes_polynomial.py
# VERSI FINAL: Memperbaiki indentasi dan format output di web agar presisi.
import time
from colorama import Fore
from docx.shared import Pt
from .aes_reporting import (
    add_poly_text_to_paragraph, add_single_term_to_para,
    add_bold_paragraph, add_calculation_paragraph
)

def byte_to_poly_str(poly_val):
    if poly_val == 0: return "0"
    terms = []
    for i in range(poly_val.bit_length()):
        if (poly_val >> i) & 1:
            if i == 0: terms.append("1")
            elif i == 1: terms.append("x")
            else: terms.append(f"x^{i}")
    return " + ".join(reversed(terms))

def poly_multiply(poly1, poly2):
    res = 0
    for i in range(8):
        if (poly2 >> i) & 1: res ^= (poly1 << i)
    return res

def explain_gmul_poly(a, b, doc):
    poly_a_str = byte_to_poly_str(a)
    poly_b_str = byte_to_poly_str(b)

    # === BAGIAN YANG DIPERBAIKI 1: Indentasi Header ===
    header_line1 = f"{a:02X} = {a:08b} = {poly_a_str}"
    header_line2 = f"{b:02X} = {b:08b} = {poly_b_str}"
    print(f"      {Fore.CYAN}{header_line1}")
    print(f"      {Fore.CYAN}{header_line2}") # Dicetak terpisah agar indentasi konsisten
    if doc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        add_poly_text_to_paragraph(p, f"{header_line1}\n{header_line2}")

    raw_product = poly_multiply(a, b)
    # === BAGIAN YANG DIPERBAIKI 2: Tanda Kurung dan Indentasi ===
    mult_result_text = f"      ({a:02X} * {b:02X}) = {byte_to_poly_str(raw_product)}"
    print(mult_result_text)
    if doc:
        p_mult = doc.add_paragraph(); p_mult.paragraph_format.left_indent = Pt(18)
        # Hapus indentasi dari string agar bisa diatur oleh paragraf
        add_poly_text_to_paragraph(p_mult, mult_result_text.strip())

    if raw_product < 0x100:
        final_result = raw_product
    else:
        reducer_poly = 0x1B
        highest_power_term = 1 << (raw_product.bit_length() - 1)
        remaining_terms_poly = raw_product ^ highest_power_term
        final_result = remaining_terms_poly ^ reducer_poly

        info_text = f"{byte_to_poly_str(highest_power_term)} is reduced by {byte_to_poly_str(reducer_poly)}"
        print(f"         {Fore.YELLOW}{info_text}")
        if doc: p_info = doc.add_paragraph(); p_info.add_run(info_text).italic = True; p_info.paragraph_format.left_indent = Pt(36)

        terms1_str = byte_to_poly_str(remaining_terms_poly)
        terms2_str = byte_to_poly_str(reducer_poly)
        
        console_display_str = f"{terms1_str} + {terms2_str}"
        print(f"         {Fore.YELLOW}{console_display_str}")

        if doc:
            p_strikethrough = doc.add_paragraph()
            p_strikethrough.paragraph_format.left_indent = Pt(36)
            terms1 = terms1_str.split(' + ')
            terms2 = terms2_str.split(' + ')
            duplicates = set(terms1) & set(terms2)

            for i, term in enumerate(terms1):
                add_single_term_to_para(p_strikethrough, term, term in duplicates)
                if i < len(terms1) - 1: p_strikethrough.add_run(' + ')
            p_strikethrough.add_run(' + ')
            for i, term in enumerate(terms2):
                add_single_term_to_para(p_strikethrough, term, term in duplicates)
                if i < len(terms2) - 1: p_strikethrough.add_run(' + ')

    # === BAGIAN YANG DIPERBAIKI 3: Indentasi Hasil ===
    final_text = f"= {byte_to_poly_str(final_result)}"
    final_text_full = f"= {byte_to_poly_str(final_result)} (Bin: {final_result:08b}) (Hex: {final_result:02X})"
    
    print(f"         {final_text}")
    print(f"         {Fore.GREEN}{final_text_full}\n")
    if doc:
        p_final = doc.add_paragraph(); p_final.paragraph_format.left_indent = Pt(36)
        add_poly_text_to_paragraph(p_final, final_text)
        p_full = doc.add_paragraph(); p_full.paragraph_format.left_indent = Pt(36)
        add_poly_text_to_paragraph(p_full, final_text_full)

    time.sleep(0.01)
    return final_result