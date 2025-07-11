# aes_reporting.py
# VERSI DEFINITIF FINAL 2: Memperbaiki fitur coret dan menyeragamkan indentasi.
import re
from colorama import Fore
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_bold_paragraph(doc, text):
    if doc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18) # Menambahkan indentasi
        p.add_run(text).bold = True

def add_calculation_paragraph(doc, text, indent=True): # Default indent sekarang True
    if doc:
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(text.strip())
        font = run.font; font.name = 'Times New Roman'; font.size = Pt(12)

def add_poly_text_to_paragraph(p, text):
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        
        parts = re.split(r'(x\^\d+)', line)
        for part in parts:
            if part.startswith('x^'):
                base, exponent = part.split('^')
                run = p.add_run(base)
                run.font.name = 'Times New Roman'; run.font.size = Pt(12)
                run_sup = p.add_run(exponent)
                run_sup.font.superscript = True
                run_sup.font.name = 'Times New Roman'; run_sup.font.size = Pt(12)
            else:
                run = p.add_run(part)
                run.font.name = 'Times New Roman'; run.font.size = Pt(12)

def add_single_term_to_para(p, term, is_duplicate):
    font_attrs = {'name': 'Times New Roman', 'size': Pt(12)}
    
    if "^" in term:
        base, exponent = term.split('^', 1)
        run = p.add_run(base)
        run.font.name, run.font.size = font_attrs['name'], font_attrs['size']
        run.font.strike = is_duplicate # Terapkan coret pada basis
        run.font.bold = is_duplicate

        run_sup = p.add_run(exponent)
        run_sup.font.superscript = True
        run_sup.font.name, run_sup.font.size = font_attrs['name'], font_attrs['size']
        run_sup.font.strike = is_duplicate # Terapkan coret pada pangkat
        run_sup.font.bold = is_duplicate
    else:
        run = p.add_run(term)
        run.font.name, run.font.size = font_attrs['name'], font_attrs['size']
        run.font.strike = is_duplicate
        run.font.bold = is_duplicate

def add_matrix_to_doc(doc, label, hex_string):
    if doc:
        p = doc.add_paragraph(); p.add_run(label).italic = True
        table = doc.add_table(rows=4, cols=4); table.style = 'Table Grid'
        bytes_array = [hex_string[i:i+2] for i in range(0, len(hex_string), 2)]
        for r in range(4):
            for c in range(4):
                cell = table.cell(r, c); run = cell.paragraphs[0].add_run(bytes_array[c*4+r].upper()); run.font.name = 'Times New Roman'; run.font.size = Pt(12); cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_constant_matrix_to_doc(doc, label, matrix):
    if doc:
        p = doc.add_paragraph(); p.add_run(label).italic = True
        table = doc.add_table(rows=4, cols=4); table.style = 'Table Grid'
        for r in range(4):
            for c in range(4):
                cell = table.cell(r, c)
                run = cell.paragraphs[0].add_run(f"{matrix[r][c]:02X}")
                run.font.name = 'Times New Roman'; run.font.size = Pt(12); cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def print_matrix(label, hex_string):
    print(f"\n{Fore.CYAN + label}"); bytes_array = [hex_string[i:i+2] for i in range(0, len(hex_string), 2)];
    for r in range(4): print(f"| {' | '.join(bytes_array[c*4+r] for c in range(4)).upper()} |")

def print_constant_matrix(label, matrix):
    print(f"\n{Fore.CYAN + label}")
    for r in range(4):
        print(f"| {' | '.join(f'{val:02X}' for val in matrix[r])} |")

def print_interim_matrix(label, data_bytes, cols_filled):
    print(f"\n{Fore.CYAN + label}");
    for r in range(4): items = [f"{data_bytes[c*4+r]:02X}" if c<cols_filled else "XX" for c in range(4)]; print(f"| {' | '.join(items)} |")