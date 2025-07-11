# interactive_aes_tester.py
import time
from datetime import date, datetime
import re

# --- Inisialisasi Pustaka ---
try:
    from colorama import Fore, Style, init
    init(autoreset=True)
except ImportError:
    print("Peringatan: library 'colorama' tidak ditemukan. Output tidak akan berwarna.")
    class AnsiFore:
        YELLOW = GREEN = MAGENTA = CYAN = BLUE = RED = WHITE = ""
    class AnsiStyle:
        RESET_ALL = ""
    Fore = AnsiFore
    Style = AnsiStyle

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = Pt = WD_ALIGN_PARAGRAPH = RGBColor = None

# --- Impor semua fungsi 'explain' dari modul yang ada ---
from aes_key_schedule import key_schedule_explain
from aes_sub_bytes import sub_bytes_explain
from aes_shift_rows import shift_rows_explain
from aes_mix_columns import mix_columns_poly_explain
from aes_add_round_key import add_round_key_explain
from aes_reporting import print_matrix, add_matrix_to_doc

# --- Fungsi Bantuan untuk Validasi Input ---
def get_validated_hex_input(prompt, length=32):
    while True:
        user_input = input(prompt).strip()
        if len(user_input) != length:
            print(f"{Fore.RED}Error: Input harus tepat {length} karakter heksadesimal.")
            continue
        try:
            bytearray.fromhex(user_input)
            return user_input
        except ValueError:
            print(f"{Fore.RED}Error: Input mengandung karakter yang bukan heksadesimal (0-9, a-f).")

def create_report_document(title):
    if not DOCX_AVAILABLE or docx is None:
        return None
    doc = docx.Document()
    style = doc.styles['Normal']
    font = getattr(style, 'font', None)
    if font is not None:
        font.name = 'Times New Roman'
        if Pt is not None:
            font.size = Pt(12)
    doc.add_heading(title, level=1)
    return doc

# --- Alur Proses Enkripsi Penuh (dari file asli) ---
def aes_full_process():
    if not DOCX_AVAILABLE:
        print(f"{Fore.RED}ERROR: Library 'python-docx' tidak ditemukan.")
        return
    try:
        plaintext_txt = input("Enter Plaintext (text 16 characters): ")
        if len(plaintext_txt) != 16: raise ValueError("Plaintext must be exactly 16 characters")
        kunci_txt = input("Enter Key (text 16 characters): ")
        if len(kunci_txt) != 16: raise ValueError("Key must be exactly 16 characters")

        nama_file_laporan = f'Enkripsi_AES_{plaintext_txt.replace(" ", "_")}_{datetime.now().strftime("%Y-%m-%d_%H-%M-%S")}.docx'
        doc = create_report_document(f'Enkripsi AES {plaintext_txt}')
        
        plaintext_hex, kunci_hex = plaintext_txt.encode('windows-1252').hex(), kunci_txt.encode('windows-1252').hex()

        if doc:
            p = doc.add_paragraph(); p.add_run('Plaintext: ').bold = True; p.add_run(f"{plaintext_txt} ({plaintext_hex.upper()})")
            p.add_run('\nKey: ').bold = True; p.add_run(f"{kunci_txt} ({kunci_hex.upper()})")

        print(f"\n{Fore.CYAN}--- Verifikasi Input ---")
        print(f"Plaintext '{plaintext_txt}' dikonversi menjadi Heks: {Fore.YELLOW}{plaintext_hex.upper()}")
        print(f"Key '{kunci_txt}' dikonversi menjadi Heks: {Fore.YELLOW}{kunci_hex.upper()}")
        
        round_keys = key_schedule_explain(kunci_hex, doc)
        
        print(f"\n{Fore.YELLOW}===== Key Schedule Result =====")
        if doc: doc.add_heading("Key Schedule Result", level=2)
        for i, key_hex in enumerate(round_keys):
            print_matrix(f"Round Key {i}", key_hex); add_matrix_to_doc(doc, f"Round Key {i}", key_hex)
        
        print(f"\n{Fore.YELLOW}===== Process Encryption =====")
        if doc: doc.add_heading("Process Encryption", level=2)
        state = plaintext_hex

        if doc: doc.add_heading("Initial Round (Pre-Round)", level=2)
        state = add_round_key_explain(state, round_keys[0], 0, doc)
        
        for i in range(1, 10):
            print(f"\n{Fore.YELLOW}===== RONDE {i} =====")
            if doc: doc.add_heading(f"RONDE {i}", level=2)
            state = sub_bytes_explain(state, doc)
            state = shift_rows_explain(state, doc)
            state = mix_columns_poly_explain(state, doc)
            state = add_round_key_explain(state, round_keys[i], i, doc)
            
        print(f"\n{Fore.YELLOW}===== RONDE 10 (Final) =====")
        if doc: doc.add_heading("RONDE 10 (Final)", level=2)
        state = sub_bytes_explain(state, doc)
        state = shift_rows_explain(state, doc)
        text_skip = "Steps: MixColumns (SKIPPED IN THE FINAL ROUND)"; print(f"\n{Fore.MAGENTA}--- {text_skip} ---")
        if doc: doc.add_heading(text_skip, level=3)
        state = add_round_key_explain(state, round_keys[10], 10, doc)

        print(f"\n{Fore.GREEN}===== PROSES SELESAI =====")
        if doc: doc.add_heading("Final Result", level=1)
        print_matrix("Final Ciphertext", state); add_matrix_to_doc(doc, "Final Ciphertext", state)
        
        print(f"\n{Fore.CYAN}Ciphertext (Hex): {Fore.YELLOW}{state.upper()}")
        if doc:
            p = doc.add_paragraph(); p.add_run('Ciphertext: ').bold = True
            p.add_run(state.upper())
            
        if doc is not None:
            doc.save(nama_file_laporan)
            print(f"\n{Fore.GREEN}Laporan berhasil dibuat dan disimpan sebagai '{nama_file_laporan}'")
        else:
            print(f"\n{Fore.RED}Laporan tidak dapat dibuat karena library 'python-docx' tidak tersedia.")

    except ValueError as e: print(f"\n{Fore.RED}Error: {e}")
    except Exception as e: print(f"\n{Fore.RED}Terjadi kesalahan tak terduga: {e}")

# --- Fungsi untuk Menu Interaktif ---
def main_menu():
    while True:
        print("\n" + "="*50)
        print(f"{Fore.YELLOW}      PROGRAM AES STEP-BY-STEP & TESTER")
        print("="*50)
        print("1. Uji AddRoundKey")
        print("2. Uji SubBytes")
        print("3. Uji ShiftRows")
        print("4. Uji MixColumns")
        print("5. Uji Key Schedule")
        print("6. Jalankan Enkripsi AES Penuh (seperti file asli)")
        print("7. Keluar")
        print("="*50)
        
        choice = input("Masukkan pilihan Anda (1-7): ")
        
        if choice == '6':
            aes_full_process()
        elif choice == '7':
            print("Terima kasih!"); break
        elif choice in ['1', '2', '3', '4', '5']:
            if not DOCX_AVAILABLE: print(f"{Fore.RED}Peringatan: Laporan .docx tidak akan dibuat.")
            
            if choice == '1': # AddRoundKey
                doc = create_report_document("Laporan Uji - AddRoundKey")
                state = get_validated_hex_input("Masukkan State (32 char hex): ")
                key = get_validated_hex_input("Masukkan Round Key (32 char hex): ")
                add_round_key_explain(state, key, 0, doc)
                if doc: doc.save(f"Laporan_Test_AddRoundKey({datetime.now().strftime('%Y%m%d_%H%M%S')}).docx")

            elif choice == '2': # SubBytes
                doc = create_report_document("Laporan Uji - SubBytes")
                state = get_validated_hex_input("Masukkan State (32 char hex): ")
                sub_bytes_explain(state, doc)
                if doc: doc.save(f"Laporan_Test_SubBytes({datetime.now().strftime('%Y%m%d_%H%M%S')}).docx")

            elif choice == '3': # ShiftRows
                doc = create_report_document("Laporan Uji - ShiftRows")
                state = get_validated_hex_input("Masukkan State (32 char hex): ")
                shift_rows_explain(state, doc)
                if doc: doc.save(f"Laporan_Test_ShiftRows({datetime.now().strftime('%Y%m%d_%H%M%S')}).docx")

            elif choice == '4': # MixColumns
                doc = create_report_document("Laporan Uji - MixColumns")
                state = get_validated_hex_input("Masukkan State (32 char hex): ")
                mix_columns_poly_explain(state, doc)
                if doc: doc.save(f"Laporan_Test_MixColumns({datetime.now().strftime('%Y%m%d_%H%M%S')}).docx")

            elif choice == '5': # Key Schedule
                doc = create_report_document("Laporan Uji - Key Schedule")
                key = get_validated_hex_input("Masukkan Kunci Awal (32 char hex): ")
                key_schedule_explain(key, doc)
                if doc: doc.save(f"Laporan_Test_KeySchedule({datetime.now().strftime('%Y%m%d_%H%M%S')}).docx")

            if DOCX_AVAILABLE: print(f"\n{Fore.GREEN}Laporan uji berhasil disimpan.")
        else:
            print(f"{Fore.RED}Pilihan tidak valid.")
        
        input(f"\n{Fore.CYAN}Tekan Enter untuk kembali ke menu...")

if __name__ == "__main__":
    main_menu()