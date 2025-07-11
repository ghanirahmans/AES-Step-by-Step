# main.py
# File utama untuk menjalankan proses enkripsi AES.
import time
from datetime import datetime

# --- Inisialisasi Pustaka ---
try:
    from colorama import Fore, Style, init
    init(autoreset=True)
except ImportError:
    print("Warning: 'colorama' library not found. Output will not be colored.")
    class Fore: YELLOW = GREEN = MAGENTA = CYAN = BLUE = RED = WHITE = ""
    class Style: RESET_ALL = ""

try:
    import docx
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = Pt = RGBColor = None

# --- Impor Modul-modul AES ---
from aes_key_schedule import key_schedule_explain
from aes_sub_bytes import sub_bytes_explain
from aes_shift_rows import shift_rows_explain
from aes_mix_columns import mix_columns_poly_explain
from aes_add_round_key import add_round_key_explain
from aes_reporting import print_matrix, add_matrix_to_doc

def aes_full_process():
    if not DOCX_AVAILABLE:
        print(f"{Fore.RED}ERROR: Library 'python-docx' not found.\n{Fore.YELLOW}Please install by running: pip install python-docx")
        return
    try:
        plaintext_txt = input("Enter Plaintext (16 characters): ")
        if len(plaintext_txt) != 16: raise ValueError("Plaintext must be exactly 16 characters")
        kunci_txt = input("Enter Key (16 characters): ")
        if len(kunci_txt) != 16: raise ValueError("Key must be exactly 16 characters")

        today_str = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        nama_file_laporan = f'Enkripsi_AES_{plaintext_txt.replace(" ", "_")}_{today_str}.docx'
        
        doc = docx.Document()
        style = doc.styles['Normal']; font = style.font; font.name = 'Times New Roman'; font.size = Pt(12); font.color.rgb = RGBColor(0,0,0)

        plaintext_hex, kunci_hex = plaintext_txt.encode('windows-1252').hex(), kunci_txt.encode('windows-1252').hex()

        doc.add_heading(f'AES Encryption for "{plaintext_txt}"', level=1)
        p = doc.add_paragraph(); p.add_run('Plaintext: ').bold = True; p.add_run(f"{plaintext_txt} ({plaintext_hex.upper()})")
        p.add_run('\nKey: ').bold = True; p.add_run(f"{kunci_txt} ({kunci_hex.upper()})")

        print(f"\n{Fore.CYAN}--- Input Verification ---")
        print(f"Plaintext '{plaintext_txt}' to Hex: {Fore.YELLOW}{plaintext_hex.upper()}")
        print(f"Key '{kunci_txt}' to Hex: {Fore.YELLOW}{kunci_hex.upper()}")
        time.sleep(1)
        
        round_keys = key_schedule_explain(kunci_hex, doc)
        
        doc.add_page_break()
        doc.add_heading("Encryption Process", level=2)
        print(f"\n{Fore.YELLOW}===== Encryption Process =====")
        state = plaintext_hex

        doc.add_heading("Initial Round (Pre-Round)", level=2)
        state = add_round_key_explain(state, round_keys[0], 0, doc)
        
        for i in range(1, 10):
            doc.add_page_break()
            doc.add_heading(f"ROUND {i}", level=2)
            print(f"\n{Fore.YELLOW}===== ROUND {i} =====")
            state = sub_bytes_explain(state, doc)
            state = shift_rows_explain(state, doc)
            state = mix_columns_poly_explain(state, doc)
            state = add_round_key_explain(state, round_keys[i], i, doc)
            
        doc.add_page_break()
        doc.add_heading("ROUND 10 (Final)", level=2)
        print(f"\n{Fore.YELLOW}===== ROUND 10 (Final) =====")
        state = sub_bytes_explain(state, doc)
        state = shift_rows_explain(state, doc)
        text_skip = "Steps: MixColumns (SKIPPED IN THE FINAL ROUND)"; print(f"\n{Fore.MAGENTA}--- {text_skip} ---")
        if doc: doc.add_heading(text_skip, level=3)
        state = add_round_key_explain(state, round_keys[10], 10, doc)

        doc.add_page_break()
        doc.add_heading("Final Result", level=1)
        print(f"\n{Fore.GREEN}===== PROCESS COMPLETE =====")
        print_matrix("Final Ciphertext", state); add_matrix_to_doc(doc, "Final Ciphertext", state)
        print(f"\n{Style.RESET_ALL}Ciphertext (Hex): {Fore.YELLOW}{state.upper()}")
        
        if doc:
            p = doc.add_paragraph(); p.add_run('Final Ciphertext (Hex): ').bold = True
            p.add_run(state.upper())
            
        doc.save(nama_file_laporan)
        print(f"\n{Fore.GREEN}Report successfully created as '{nama_file_laporan}'")

    except ValueError as e: print(f"\n{Fore.RED}Error: {e}")
    except Exception as e: print(f"\n{Fore.RED}An unexpected error occurred: {e}")

if __name__ == "__main__":
    aes_full_process()